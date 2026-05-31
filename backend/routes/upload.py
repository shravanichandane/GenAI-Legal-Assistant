### backend/routes/upload.py

from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy.orm import Session
from ..database import get_db
from ..models import Document, DocumentCreate, DocumentResponse
from ..services.clause_extractor import clause_extractor
import PyPDF2
import docx
import io
import logging
from ..auth import get_current_user, User

logger = logging.getLogger(__name__)
router = APIRouter()

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Validate file type
    allowed_extensions = ('.pdf', '.docx', '.txt')
    if not file.filename.lower().endswith(allowed_extensions):
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed formats: {', '.join(allowed_extensions)}"
        )
    
    # Validate file size (max 10MB)
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 10MB.")
    
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")
    
    # Extract text based on file type
    try:
        text_content = extract_text_from_file(content, file.filename)
        
        if not text_content or len(text_content.strip()) < 10:
            raise HTTPException(status_code=400, detail="No readable text found in the document.")
            
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")
    
    # Create document record
    try:
        document_data = DocumentCreate(filename=file.filename, content=text_content)
        db_document = Document(**document_data.dict(), user_id=current_user.id)
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        logger.info(f"Successfully uploaded document: {file.filename} (ID: {db_document.id})")
        return db_document
        
    except Exception as e:
        db.rollback()
        logger.error(f"Database error uploading {file.filename}: {e}")
        raise HTTPException(status_code=500, detail="Failed to save document to database.")

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract text from uploaded file based on type"""
    if filename.lower().endswith('.pdf'):
        return extract_pdf_text(content)
    elif filename.lower().endswith('.docx'):
        return extract_docx_text(content)
    elif filename.lower().endswith('.txt'):
        try:
            # Try UTF-8 first, then fallback to other encodings
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Failed to decode text file: {str(e)}")
    else:
        raise ValueError("Unsupported file type")

def extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF"""
    try:
        pdf_file = io.BytesIO(content)
        reader = PyPDF2.PdfReader(pdf_file)
        
        if len(reader.pages) == 0:
            raise ValueError("PDF contains no pages")
        
        text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:  # Only add if text was extracted
                    text += page_text + "\n"
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                continue
        
        if not text.strip():
            raise ValueError("No text could be extracted from PDF - document may be image-based or corrupted")
            
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX"""
    try:
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        if not text_parts:
            raise ValueError("No text content found in DOCX file")
        
        return "\n".join(text_parts)
        
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
